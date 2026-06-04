#!/usr/bin/env python3
"""Generate Arabic Company Profile Word Document"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Set paragraph to RTL (Right-to-Left) for Arabic"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def create_word_doc():
    doc = Document()
    
    # Set default font for Arabic
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    
    # Colors
    navy = RGBColor(0x1E, 0x3A, 0x5F)
    brown = RGBColor(0x8B, 0x45, 0x13)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title)
    run = title.add_run("الملف التعريفي للشركة")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = navy
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(subtitle)
    run = subtitle.add_run("شركة أركان الراي (ARETION)")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = navy
    
    doc.add_paragraph()
    
    # Helper function to add heading
    def add_heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = navy
        return p
    
    # Helper function to add subheading
    def add_subheading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = brown
        return p
    
    # Helper function to add body text
    def add_body(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        run = p.add_run(text)
        run.font.size = Pt(12)
        return p
    
    # Helper function to add bullet
    def add_bullet(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(p)
        run = p.add_run(f"• {text}")
        run.font.size = Pt(12)
        return p
    
    # =====================
    # CONTENT
    # =====================
    
    # Company Overview
    add_heading("نبذة تعريفية عن الشركة")
    add_body("شركة أركان الراي (ARETION) هي شركة تقنية سعودية متخصصة في تصميم وتطوير حلول تقنية مبتكرة ومتقدمة في مجالات إدارة الكوارث وحماية البنية التحتية الحيوية. تعمل الشركة على تحويل أنظمة السلامة والاستجابة للطوارئ من خلال منصات ذكية مدعومة بالذكاء الاصطناعي.")
    add_body("تقدم الشركة تسعة حلول متكاملة تغطي الإنذار المبكر، والتحليلات التنبؤية، والرعاية الصحية عن بُعد، وإدارة استمرارية الأعمال، مما يجعلها الشريك الأمثل للجهات الحكومية والبلديات والمنشآت الصحية.")
    
    doc.add_paragraph()
    
    # Vision and Mission
    add_heading("الرؤية والرسالة")
    add_subheading("الرؤية")
    add_body("أن نكون الشريك التقني الرائد في تحويل منظومة السلامة للبنية التحتية الحيوية على مستوى المنطقة.")
    add_subheading("الرسالة")
    add_body("تحويل سلامة البنية التحتية الحيوية من خلال التقنيات المبتكرة والحلول الذكية التي تتنبأ بالمخاطر قبل وقوعها وتستجيب لها بشكل آلي وفوري.")
    
    doc.add_paragraph()
    
    # Solutions
    add_heading("الحلول والخدمات المقدمة")
    add_body("تقدم شركة أركان الراي تسعة حلول تقنية متكاملة:")
    
    solutions = [
        ("1. نظام إدارة الكوارث (DisasterMS)", "منصة شاملة للاستجابة للكوارث مدعومة بالذكاء الاصطناعي، توفر أنظمة إنذار مبكر واستجابة آلية متكاملة."),
        ("2. المساعد الذكي (AI Assistant)", "محرك ذكاء اصطناعي متطور بتقنية RAG يوفر دقة عالية في الوصول للمعرفة المؤسسية واتخاذ القرارات."),
        ("3. التحليلات التنبؤية (Predictive Analytics)", "نماذج تعلم آلي متقدمة للتنبؤ بالمخاطر المناخية والكوارث الطبيعية قبل وقوعها."),
        ("4. نظام الفرز الجماعي (Mass Triage System)", "تقنية رقمية مبتكرة لتحديد الهوية والفرز في حالات الكوارث الجماعية."),
        ("5. شبكة الطوارئ الطبية (EM:CC Network)", "شبكة متكاملة تربط المنشآت الصحية الإقليمية لإدارة رحلة المريض بسلاسة."),
        ("6. التنبيب عن بُعد (Tele-Intubation)", "حل روبوتي للتنبيب عن بُعد يوفر خبرة متخصصة في إدارة مجرى الهواء الحرج."),
        ("7. حقيبة الشفرة الزرقاء (Code Blue Kit)", "نظام متكامل من الأجهزة والبرمجيات لتنسيق حالات الطوارئ عن بُعد."),
        ("8. مصمم البروتوكولات (Protocol Designer)", "منصة مدعومة بالذكاء الاصطناعي لتصميم وتنفيذ بروتوكولات الطوارئ."),
        ("9. منصة الاستشارات (Consultation Platform)", "بنية تحتية مرنة للرعاية الصحية عن بُعد تخدم قطاعات الطب عن بُعد والاستشارات والتعليم."),
    ]
    
    for title, desc in solutions:
        add_subheading(title)
        add_body(desc)
    
    doc.add_page_break()
    
    # Target Sectors
    add_heading("القطاعات والعملاء المستهدفون")
    add_body("تخدم شركة أركان الراي القطاعات التالية:")
    add_bullet("البلديات: حلول إدارة الكوارث والطوارئ للمدن والمناطق")
    add_bullet("الوزارات والجهات الحكومية: أنظمة حماية البنية التحتية الحيوية")
    add_bullet("القطاع الصحي: منصات الرعاية الصحية عن بُعد وإدارة الطوارئ الطبية")
    add_bullet("المنشآت الحيوية: حلول الإنذار المبكر واستمرارية الأعمال")
    
    doc.add_paragraph()
    
    # Competitive Advantages
    add_heading("المزايا التنافسية")
    advantages = [
        "تقنيات مبتكرة: حلول تقنية متقدمة تتحدى الوضع الراهن وتفتح آفاقاً جديدة في مجال السلامة",
        "التركيز على إدارة الكوارث: حلول مصممة خصيصاً للاستجابة للطوارئ وإدارة الأزمات",
        "الذكاء الاصطناعي أولاً: التعلم الآلي والأتمتة في صميم كل حل نقدمه",
        "ابتكار مثبت: أربع براءات اختراع قيد التسجيل تحمي منهجياتنا المبتكرة",
        "بنية قابلة للتوسع: منصات سحابية مصممة للنشر السريع والنمو المستدام",
        "حلول مختبرة ميدانياً: تم تطويرها واختبارها من قِبل متخصصين في حالات طوارئ حقيقية",
        "جاهزية التكامل: اتصال سلس مع البنية التحتية والأنظمة القائمة",
        "تقنيات مستقبلية: تطور مستمر للبقاء في طليعة التحديات الناشئة",
        "موثوقية حرجة: أنظمة مصممة لضمان عدم التوقف عند الحاجة الماسة",
    ]
    for adv in advantages:
        add_bullet(adv)
    
    doc.add_paragraph()
    
    # Certifications
    add_heading("الاعتمادات والشهادات")
    add_body("تحرص شركة أركان الراي على الالتزام بأعلى معايير الأمان والجودة:")
    certs = [
        "CSA STAR Level Two: برنامج ضمان أمان السحابة من CSA (تدقيق طرف ثالث)",
        "ISO/IEC 27001: معيار إدارة أمن المعلومات",
        "ISO/IEC 27017: ممارسات أمان المعلومات الخاصة بالسحابة",
        "ISO/IEC 27018: حماية البيانات الشخصية في السحابة",
        "PCI DSS v4.0.1: معيار أمان بيانات صناعة بطاقات الدفع",
    ]
    for cert in certs:
        add_bullet(cert)
    
    doc.add_paragraph()
    
    # Market Opportunity
    add_heading("الفرص السوقية")
    add_body("تعمل الشركة في أسواق واعدة ومتنامية بإجمالي يتجاوز 500 مليار دولار:")
    markets = [
        "حماية البنية التحتية الحيوية: 197 مليار دولار (نمو سنوي 5.1%)",
        "أنظمة الاستعداد للكوارث: 308 مليار دولار (نمو سنوي 8.2%)",
        "كشف الكوارث الطبيعية (إنترنت الأشياء): 25.2 مليار دولار (نمو سنوي 36.3%)",
        "إدارة استمرارية الأعمال: 2.09 مليار دولار (نمو سنوي 15.5%)",
    ]
    for market in markets:
        add_bullet(market)
    
    doc.add_page_break()
    
    # Business Model
    add_heading("نموذج الأعمال")
    add_body("تعتمد الشركة على مصادر إيرادات متنوعة ومستدامة:")
    revenue = [
        "الأجهزة والمعدات: أجهزة الاستشعار والمراقبة مع عقود صيانة",
        "منصة SaaS: اشتراكات سنوية للمراقبة والتحليلات السحابية",
        "الاستشارات والتنفيذ: خدمات النشر والتكامل والتدريب",
        "العقود الحكومية: اتفاقيات إطارية متعددة السنوات",
        "البيانات والذكاء: تحليلات تنبؤية متقدمة وتقييم المخاطر",
        "ترخيص التقنية: ترخيص الملكية الفكرية للشركاء والموزعين",
    ]
    for item in revenue:
        add_bullet(item)
    
    doc.add_paragraph()
    
    # Golden Minute Achievement
    add_heading("إنجاز بارز: الدقيقة الذهبية")
    add_body("نفخر في شركة أركان الراي بتحقيق إنجاز استثنائي في تحسين زمن الاستجابة للطوارئ، حيث نجحنا في تقليص مفهوم \"الساعة الذهبية\" التقليدي إلى الدقيقة الذهبية من خلال أنظمة الإنذار المبكر والاستجابة الآلية المدعومة بالذكاء الاصطناعي.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Contact Information
    add_heading("معلومات التواصل")
    
    contact_lines = [
        "شركة أركان الراي",
        "مسجلة في المملكة العربية السعودية",
        "",
        "العنوان:",
        "مركز الملك عبدالله المالي (كافد)",
        "شارع الابتكار، العقيق",
        "مبنى 7229، الرياض 13519",
        "المملكة العربية السعودية",
        "",
        "الهاتف: 6458 525 11 966+",
        "البريد الإلكتروني: Contact@aretion.org",
        "الموقع الإلكتروني: https://aretion.org",
    ]
    
    for line in contact_lines:
        if line:
            add_body(line)
        else:
            doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p)
    run = p.add_run("✓ جميع منتجاتنا جاهزة للنشر والتطبيق")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = navy
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(p)
    run = p.add_run("© 2026 شركة أركان الراي. جميع الحقوق محفوظة.")
    run.font.size = Pt(10)
    
    # Save document
    doc.save("/app/frontend/public/ARETION_Company_Profile_Arabic.docx")
    print("Word document created successfully!")

if __name__ == "__main__":
    create_word_doc()
